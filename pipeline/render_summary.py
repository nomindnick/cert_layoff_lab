#!/usr/bin/env python3
"""Stage 4: render an annual Layoff Decision Summaries report from the merged
decision records — the deterministic "button push" that regenerates the
traditional volume for any extracted year.

The report mirrors the human volumes' format: Roman-numeral issue sections in
the taxonomy's canonical order, one-paragraph holdings each ending in the
"District (ALJ)" citation convention, and an appendix listing every decision
in the dataset by OAH case number. Entry text is the extraction's
`summary_style_holding` (verbatim; written for exactly this register), so the
document is an assembly over the corpus, not a synthesis pass — every entry
traces to quote-anchored structured data in its decision record.

Outputs (under output/reports/{year}/):
  Layoff_Decision_Summaries_{year}_AI.docx   the report
  entries_{year}.json                        section/entry manifest, reused by
                                             annotate_summary.py so the
                                             annotated HTML renders the same
                                             entries in the same order

Usage: render_summary.py --year 2009
"""

import argparse
import datetime
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DECISIONS = ROOT / "output" / "corpus" / "decisions"
TAXONOMY = ROOT / "output" / "summaries" / "taxonomy.json"

# Section headings in the human volumes' register. Categories absent from the
# map fall back to the canonical id, upper-cased.
DISPLAY = {
    "procedural_issues": "PROCEDURAL ISSUES",
    "calculations_ada_fte": "ADA AND FTE CALCULATIONS",
    "attrition": "ATTRITION",
    "pks_reduction": "REDUCTION OF PARTICULAR KINDS OF SERVICES",
    "seniority": "ISSUES RELATED TO SENIORITY",
    "temporary_employees": "TEMPORARY EMPLOYEES",
    "categorically_funded": "CATEGORICALLY FUNDED PROGRAMS",
    "substitutes": "SUBSTITUTE EMPLOYEES",
    "skipping": "SKIPPING",
    "bumping": "BUMPING",
    "assignments_reassignments": "ASSIGNMENTS AND REASSIGNMENTS",
    "credentials": "CREDENTIALS",
    "competency": "COMPETENCY",
    "tie_breaking": "TIE-BREAKING CRITERIA",
    "domino_theory": "DOMINO THEORY",
    "eera_cba_aa": "EERA AND COLLECTIVE BARGAINING ISSUES",
    "contractual_issues": "CONTRACTUAL ISSUES",
    "county_office_issues": "COUNTY OFFICE OF EDUCATION ISSUES",
    "discrimination": "DISCRIMINATION",
    "adult_education": "ADULT EDUCATION",
    "reemployment_rights": "REEMPLOYMENT RIGHTS",
    "miscellaneous": "MISCELLANEOUS",
    "other": "OTHER ISSUES",
}

OUTCOME_DISPLAY = {
    "sustained": "Accusation sustained",
    "sustained_in_part": "Sustained in part",
    "not_sustained": "Not sustained",
}

# The schema's outcome-neutral pks_reduction sits where the volumes' outcome
# pair sat; both legacy ids collapse onto it for ordering.
LEGACY_CATEGORY = {"pks_allowed": "pks_reduction", "pks_not_allowed": "pks_reduction"}

_CITE_TAIL = re.compile(r"\([A-Z][a-zA-Z.À-ſ'\- ]+\)\.?\s*$")
_SMALL_WORDS = {"of", "and", "the", "for", "de", "del", "la", "las", "los"}
_ALJ_SUFFIXES = {"jr", "jr.", "sr", "sr.", "ii", "iii", "iv", "esq", "esq."}


def roman(n):
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
            (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
            (5, "V"), (4, "IV"), (1, "I")]
    out = []
    for v, s in vals:
        while n >= v:
            out.append(s)
            n -= v
    return "".join(out)


def title_case(s):
    words = (s or "").strip().split()
    out = []
    for i, w in enumerate(words):
        lw = w.lower()
        out.append(lw if (lw in _SMALL_WORDS and i > 0) else lw.capitalize())
    return " ".join(out)


def district_short(raw):
    """ "SAN JUAN UNIFIED SCHOOL DISTRICT" -> "San Juan Unified" — the
    volumes' citation style keeps the organizational qualifier (Unified,
    Union High, ...) but drops the "School District" tail."""
    t = title_case(raw)
    t = re.sub(r"\s+School District$", "", t)
    t = re.sub(r"\s+District$", "", t)
    return t


def alj_surname(raw):
    toks = [t for t in (raw or "").replace(",", " ").split()
            if t.lower() not in _ALJ_SUFFIXES]
    return toks[-1] if toks else ""


def load_decisions(year):
    recs = {}
    for f in sorted(DECISIONS.glob(f"{year}*.json")):
        recs[f.stem] = json.loads(f.read_text())
    return recs


def category_order():
    tax = json.loads(TAXONOMY.read_text())
    order, seen = [], set()
    for c in tax["canonical_order"]:
        c = LEGACY_CATEGORY.get(c, c)
        if c not in seen:
            order.append(c)
            seen.add(c)
    return order


def deidentify(text, rec):
    """Replace respondent names with their pseudonymous roster refs (R1..Rn).

    Respondent names are private-layer only (CLAUDE.md); the report follows
    the human volumes' convention of never naming teachers. Substitution is
    deterministic from the record's own roster: full name first (longest
    match), then bare capitalized surname — guarded against colliding with
    the ALJ's surname or the district name, where the same string is a
    different person. Over-redaction is the safe direction; under-redaction
    is a privacy defect, so ambiguous surnames still redact."""
    ident = rec.get("identity") or {}
    alj = alj_surname((ident.get("alj") or {}).get("raw") or "").lower()
    dist = ((ident.get("district") or {}).get("raw") or "").lower()
    subs = []  # (match string, ref)
    firsts = set()  # roster first-name tokens, for the stranded-name pass
    for r in (rec.get("outcome") or {}).get("roster") or []:
        name, ref = (r.get("name") or "").strip(), r.get("ref")
        if not name or not ref:
            continue
        if "," in name:  # "Last, First" form
            last = name.split(",")[0].strip()
            first = name.split(",", 1)[1].strip()
            subs.append((f"{first} {last}", ref))
        else:
            parts = name.split()
            last, first = parts[-1], parts[0]
        subs.append((name, ref))
        firsts.add(first.split()[0].lower())
        if (len(last) >= 3 and last.lower() != alj
                and last.lower() not in dist):
            subs.append((last, ref))
    subs.sort(key=lambda s: -len(s[0]))
    n = 0
    for s, ref in subs:
        # case-insensitive: rosters are often ALL CAPS while the model writes
        # names in title case
        pat = re.compile(rf"\b{re.escape(s)}\b", re.IGNORECASE)
        text, k = pat.subn(ref, text)
        n += k
    # Fuzzy backstop: extraction sometimes spells a respondent's surname one
    # edit off the roster ("Myer" / "Myers"), which the exact pass misses.
    # A unique distance-1 surname match redacts to that ref; an ambiguous one
    # (two Myers on the roster) redacts to a neutral placeholder — privacy
    # over per-respondent fidelity.
    surnames = {}
    for s, ref in subs:
        if " " not in s:
            surnames.setdefault(s.lower(), set()).add(ref)
    for tok in set(re.findall(r"\b[A-Z][A-Za-zà-ÿ'\-]{2,}\b", text)):
        tl = tok.lower()
        if tok.isupper() or tl == alj or tl in dist or tl in _FUZZY_STOP:
            continue
        hits = set()
        for sn, refs in surnames.items():
            if abs(len(sn) - len(tl)) <= 1 and _dist1(tl, sn):
                hits |= refs
        if hits:
            repl = hits.pop() if len(hits) == 1 else "[a respondent]"
            text, k = re.subn(rf"\b{re.escape(tok)}\b", repl, text)
            n += k
    # Stranded first names: when only the surname matched (roster spelling
    # drift), the model's "Vickie Ensley" becomes "Vickie R1". Collapse a
    # leading token that is a roster first name onto the ref that follows it.
    def _strand(m):
        nonlocal n
        if m.group(1).lower() in firsts:
            n += 1
            return m.group(2)
        return m.group(0)
    text = re.sub(r"\b([A-Za-zà-ÿ'\-]+)\s+(R\d+\b|\[a respondent\])",
                  _strand, text)
    return text, n


# Frequent capitalized words in these decisions that must never fuzzy-match
# a roster surname.
_FUZZY_STOP = {
    "district", "respondent", "respondents", "education", "code", "board",
    "accusation", "april", "march", "english", "spanish", "french", "german",
    "county", "school", "services", "section", "american", "resolution",
    "february", "january", "title", "master", "doctor", "found", "order",
    "judge", "court", "state", "notice", "hearing", "evidence", "credential",
}


def _dist1(a, b):
    """Levenshtein distance <= 1 (equal, or one sub/ins/del)."""
    if a == b:
        return True
    if abs(len(a) - len(b)) > 1:
        return False
    if len(a) == len(b):
        return sum(x != y for x, y in zip(a, b)) == 1
    if len(a) > len(b):
        a, b = b, a
    for i in range(len(b)):
        if a == b[:i] + b[i + 1:]:
            return True
    return False


def entry_text(holding, cite):
    """The per-entry paragraph: summary_style_holding verbatim, with the
    District (ALJ) citation appended when the model's paragraph lacks it."""
    text = (holding.get("summary_style_holding") or "").strip()
    if not text:
        issue = (holding.get("issue") or {}).get("statement") or ""
        party = (holding.get("ruling") or {}).get("prevailing_party")
        tail = {"district": "The ALJ resolved the issue in the District's favor.",
                "respondent": "The ALJ resolved the issue in the respondent's favor.",
                "mixed": "The ALJ resolved the issue with mixed results.",
                }.get(party, "")
        text = f"{issue} {tail}".strip()
    if not _CITE_TAIL.search(text):
        text = f"{text.rstrip('.')}" + f". {cite}"
    if not text.endswith("."):
        text += "."
    return text


def build_entries(year):
    """Assemble the report structure: ordered sections of per-holding entries
    plus the decisions appendix. Pure function of the decision records."""
    recs = load_decisions(year)
    if not recs:
        sys.exit(f"no decision records under {DECISIONS} for year {year}")
    by_cat = {}
    n_holdings = 0
    total_redacted = 0
    for case_no, rec in recs.items():
        ident = rec.get("identity") or {}
        dist = district_short((ident.get("district") or {}).get("raw") or "")
        surname = alj_surname((ident.get("alj") or {}).get("raw") or "")
        cite = f"{dist} ({surname})"
        for hi, h in enumerate(rec.get("holdings") or []):
            n_holdings += 1
            cat = (h.get("issue") or {}).get("category") or "other"
            cat = LEGACY_CATEGORY.get(cat, cat)
            text, n_redacted = deidentify(entry_text(h, cite), rec)
            total_redacted = total_redacted + n_redacted
            by_cat.setdefault(cat, []).append({
                "case_no": case_no,
                "holding_idx": hi,
                "district": dist,
                "alj": surname,
                "cite": cite,
                "category": cat,
                "subtype": (h.get("issue") or {}).get("subtype"),
                "notable": bool(h.get("notable")),
                "n_redacted": n_redacted,
                "text": text,
            })
    order = [c for c in category_order() if c in by_cat]
    order += sorted(c for c in by_cat if c not in order)
    sections = []
    for i, cat in enumerate(order, 1):
        entries = sorted(by_cat[cat],
                         key=lambda e: (e["district"], e["case_no"], e["holding_idx"]))
        sections.append({
            "category": cat,
            "roman": roman(i),
            "title": DISPLAY.get(cat, cat.replace("_", " ").upper()),
            "entries": entries,
        })
    appendix = []
    for case_no, rec in sorted(recs.items()):
        ident = rec.get("identity") or {}
        out = rec.get("outcome") or {}
        appendix.append({
            "case_no": case_no,
            "district": district_short((ident.get("district") or {}).get("raw") or ""),
            "alj": alj_surname((ident.get("alj") or {}).get("raw") or ""),
            "date": ident.get("decision_date") or "",
            "outcome": OUTCOME_DISPLAY.get(out.get("overall"), out.get("overall") or ""),
        })
    return {
        "year": year,
        "n_decisions": len(recs),
        "n_holdings": n_holdings,
        "n_redacted": total_redacted,
        "sections": sections,
        "appendix": appendix,
    }


def audit_entries(data):
    """Post-render privacy audit: scan final entry text (case-insensitively)
    for any roster name or surname that survived de-identification. Runs on
    every render; a leak is loud, never silent."""
    recs = load_decisions(data["year"])
    leaks = []
    for sec in data["sections"]:
        for e in sec["entries"]:
            rec = recs[e["case_no"]]
            ident = rec.get("identity") or {}
            alj = alj_surname((ident.get("alj") or {}).get("raw") or "").lower()
            dist = ((ident.get("district") or {}).get("raw") or "").lower()
            for r in (rec.get("outcome") or {}).get("roster") or []:
                name = (r.get("name") or "").strip()
                if not name:
                    continue
                if "," in name:
                    last = name.split(",")[0].strip()
                else:
                    last = name.split()[-1]
                pats = [name]
                if len(last) >= 3 and last.lower() != alj and last.lower() not in dist:
                    pats.append(last)
                for p in pats:
                    if re.search(rf"\b{re.escape(p)}\b", e["text"], re.IGNORECASE):
                        leaks.append((e["case_no"], e["holding_idx"], p))
                        break
    return leaks


PREAMBLE = (
    "This summary was generated automatically from a structured corpus of "
    "Office of Administrative Hearings proposed decisions in certificated "
    "employee layoff proceedings under Education Code sections 44949 and "
    "44955. Each decision was processed into a structured record in which "
    "every holding, argument, and factual finding is anchored to verbatim "
    "quotations from the decision text; this document is assembled directly "
    "from those records. It covers the {n_decisions} decisions for {year} "
    "presently in the dataset ({n_holdings} catalogued holdings). Decisions "
    "not yet ingested into the dataset are not reflected. Citations follow "
    "the convention of the traditional Layoff Decision Summaries: district "
    "and administrative law judge surname, with respondents identified "
    "pseudonymously."
)


def write_docx(data, path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"SUMMARY OF {data['year']} PROPOSED ALJ LAYOFF DECISIONS")
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Generated from the structured decision corpus — "
                    f"{datetime.date.today().isoformat()}")
    r.italic = True

    doc.add_paragraph(PREAMBLE.format(**data))

    for sec in data["sections"]:
        h = doc.add_paragraph()
        hr = h.add_run(f"{sec['roman']}. {sec['title']}")
        hr.bold = True
        hr.font.size = Pt(12)
        for n, e in enumerate(sec["entries"], 1):
            p = doc.add_paragraph(f"{n}. {e['text']}")
            p.paragraph_format.space_after = Pt(8)

    h = doc.add_paragraph()
    hr = h.add_run("APPENDIX — DECISIONS IN THIS DATASET")
    hr.bold = True
    hr.font.size = Pt(12)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(["OAH Case No.", "District", "ALJ", "Decision Date",
                             "Outcome"]):
        cell = table.rows[0].cells[i]
        cell.text = col
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in data["appendix"]:
        cells = table.add_row().cells
        cells[0].text = row["case_no"]
        cells[1].text = row["district"]
        cells[2].text = row["alj"]
        cells[3].text = row["date"]
        cells[4].text = row["outcome"]
    doc.save(path)


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--year", required=True)
    args = ap.parse_args()

    data = build_entries(args.year)
    leaks = audit_entries(data)
    if leaks:
        print(f"!! PRIVACY AUDIT: {len(leaks)} roster-name leak(s) survived "
              f"de-identification — fix before sharing:")
        for case_no, hi, p in leaks[:20]:
            print(f"   {case_no} holding {hi}: {p!r}")
    else:
        print("privacy audit: no roster names in rendered entries")
    out = ROOT / "output" / "reports" / args.year
    out.mkdir(parents=True, exist_ok=True)
    manifest = out / f"entries_{args.year}.json"
    manifest.write_text(json.dumps(data, indent=1, ensure_ascii=False))
    docx_path = out / f"Layoff_Decision_Summaries_{args.year}_AI.docx"
    write_docx(data, docx_path)
    print(f"{data['n_decisions']} decisions, {data['n_holdings']} holdings, "
          f"{len(data['sections'])} sections")
    print(f"wrote {manifest}")
    print(f"wrote {docx_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
